import os
import re
import pandas as pd
import numpy as np
import sys
import datetime as dt
import docx
import comtypes.client
import shutil
from matplotlib import pyplot as plt
from matplotlib.dates import MONDAY, TUESDAY, WEDNESDAY, THURSDAY, FRIDAY, SATURDAY, SUNDAY
from matplotlib.dates import MonthLocator, WeekdayLocator, DateFormatter

# Need to change file dir once on new computer
FILE_DIR = "C:/Users/Running Injury Clini/Desktop/Marathon Training/"
PARTICIPANT_ID = str(sys.argv[1])
START_DATE = str(sys.argv[2])
END_DATE = str(sys.argv[3])


def get_dates(PARTICIPANT_ID):
    path = FILE_DIR + str(PARTICIPANT_ID)
    file_list = os.listdir(path)
    file_list = [s for s in file_list if '.' not in s]
    file_list = [s for s in file_list if re.search('[0-9]{8,}', s)]
    file_list = [s for s in file_list if int(s) >= int(START_DATE)]
    file_list = [s for s in file_list if int(s) <= int(END_DATE)]
    if len(file_list)<3:
        sys.exit("ERROR: Insufficient runs to generate report... Only {0} runs found".format(len(file_list)));
    return (file_list)

def get_participant_data(file_list):
    participant_data = []
    i = 1
    for file in file_list:
        print("{0}: Processing file {1} of {2}".format(file, i, len(file_list)))
        path1 = FILE_DIR + str(PARTICIPANT_ID) + "/" + str(file) + "/Upload"
        files = os.listdir(path1)
        date = file
        garmin_r = re.compile(".*Garmin.*")
        lumo_r = re.compile(".*Lumo.*")
        garmin_file = list(filter(garmin_r.match, files))  # Read Note
        lumo_file = list(filter(lumo_r.match, files))  # Read Note

        if len(garmin_file) > 0:
            path = path1 + "/" + str(garmin_file[0])
            garmin_data = pd.read_csv(path)
            garmin_colnames = list(garmin_data.columns.values)
            time_garmin = max(garmin_data['record.timestamp[s]'].values) - min(
                garmin_data['record.timestamp[s]'].values)
            distance = max(garmin_data['record.distance[m]'].values)
            if 'record.altitude[m]' in garmin_colnames:
                elevation_gain = max(garmin_data['record.altitude[m]'].values) - min(
                    garmin_data['record.altitude[m]'].values)
            else:
                elevation_gain = np.nan
            mean_speed = np.nanmean(garmin_data['record.speed[m/s]'].values)
            sd_speed = np.nanstd(garmin_data['record.speed[m/s]'].values)
            if 'record.heart_rate[bpm]' in garmin_colnames:
                mean_hr = np.nanmean(garmin_data['record.heart_rate[bpm]'].values)
                sd_hr = np.nanstd(garmin_data['record.heart_rate[bpm]'].values)
            elif 'record.heartrate' in garmin_colnames:
                mean_hr = np.nanmean(garmin_data['record.heartrate'].values)
                sd_hr = np.nanstd(garmin_data['record.heartrate'].values)
            elif 'record.heartrate[bpm]' in garmin_colnames:
                mean_hr = np.nanmean(garmin_data['record.heartrate[bpm]'].values)
                sd_hr = np.nanstd(garmin_data['record.heartrate[bpm]'].values)

            # mean_cadence_garmin = np.mean(garmin_data['record.cadence[rpm]'])
            # sd_cadence_garmin = np.std(garmin_data['record.cadence[rpm]'])

            garmin_input = pd.DataFrame([[time_garmin, distance, elevation_gain, mean_speed, sd_speed, mean_hr, sd_hr]],
                                        columns=['time', 'distance', 'elevation_gain', 'mean_speed', 'sd_speed',
                                                 'mean_hr',
                                                 'sd_hr'])

        if len(lumo_file) > 0:
            path = path1 + "/" + str(lumo_file[0])
            lumo_data = pd.read_csv(path)
            time_lumo = max(lumo_data['Time (s)'].values)
            mean_cadence_lumo = np.mean(lumo_data['Cadence (spm)'].values)
            sd_cadence_lumo = np.std(lumo_data['Cadence (spm)'].values)
            mean_bounce = np.mean(lumo_data['Bounce (cm)'].values)
            sd_bounce = np.std(lumo_data['Bounce (cm)'].values)
            mean_breaking = np.mean(lumo_data['Braking (m/s)'].values)
            sd_breaking = np.std(lumo_data['Braking (m/s)'].values)
            mean_drop = np.mean(lumo_data['Drop (°)'].values)
            sd_drop = np.std(lumo_data['Drop (°)'].values)
            mean_rotation = np.mean(lumo_data['Rotation (°)'].values)
            sd_rotation = np.std(lumo_data['Rotation (°)'].values)
            mean_gct = np.mean(lumo_data['GCT (ms)'].values)
            sd_gct = np.std(lumo_data['GCT (ms)'].values)

            lumo_input = pd.DataFrame(
                [[time_lumo, mean_cadence_lumo, sd_cadence_lumo, mean_bounce, sd_bounce, mean_breaking,
                  sd_breaking, mean_drop, sd_drop, mean_rotation, sd_rotation, mean_gct, sd_gct]],
                columns=['time', 'mean_cadence', 'sd_cadence', 'mean_bounce', 'sd_bounce', 'mean_breaking',
                         'sd_breaking', 'mean_drop', 'sd_drop', 'mean_rotation', 'sd_rotation', 'mean_gct', 'sd_gct'])

        if len(garmin_file) + len(lumo_file) == 2:
            input_data = [date, 'Both', garmin_input.time.values[0], garmin_input.distance.values[0],
                          garmin_input.elevation_gain.values[0],
                          garmin_input.mean_speed.values[0], garmin_input.sd_speed.values[0],
                          garmin_input.mean_hr.values[0], garmin_input.sd_hr.values[0],
                          lumo_input.mean_cadence.values[0], lumo_input.sd_cadence.values[0],
                          lumo_input.mean_bounce.values[0], lumo_input.sd_bounce.values[0],
                          lumo_input.mean_breaking.values[0], lumo_input.sd_breaking.values[0],
                          lumo_input.mean_drop.values[0], lumo_input.sd_drop.values[0],
                          lumo_input.mean_rotation.values[0], lumo_input.sd_rotation.values[0],
                          lumo_input.mean_gct.values[0], lumo_input.sd_gct.values[0]]
            participant_data.append(input_data)
        elif len(garmin_file) == 1:
            input_data = [date, 'Both', garmin_input.time.values[0], garmin_input.distance.values[0],
                          garmin_input.elevation_gain.values[0],
                          garmin_input.mean_speed.values[0], garmin_input.sd_speed.values[0],
                          garmin_input.mean_hr.values[0], garmin_input.sd_hr.values[0],
                          np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan, np.nan]
            participant_data.append(input_data)
        else:
            input_data = [date, 'Lumo', lumo_input.time.values[0], np.nan, np.nan, np.nan, np.nan, np.nan, np.nan,
                          lumo_input.mean_cadence.values[0], lumo_input.sd_cadence.values[0],
                          lumo_input.mean_bounce.values[0], lumo_input.sd_bounce.values[0],
                          lumo_input.mean_breaking.values[0], lumo_input.sd_breaking.values[0],
                          lumo_input.mean_drop.values[0], lumo_input.sd_drop.values[0],
                          lumo_input.mean_rotation.values[0], lumo_input.sd_rotation.values[0],
                          lumo_input.mean_gct.values[0], lumo_input.sd_gct.values[0]]
            participant_data.append(input_data)

        i = i + 1
        # end for loop

    output = pd.DataFrame(participant_data,
                          columns=['date', 'files', 'time', 'distance', 'elevation_gain', 'mean_speed', 'sd_speed',
                                   'mean_hr',
                                   'sd_hr', 'mean_cadence', 'sd_cadence', 'mean_bounce', 'sd_bounce', 'mean_breaking',
                                   'sd_breaking', 'mean_drop',
                                   'sd_drop', 'mean_rotation', 'sd_rotation', 'mean_gct', 'sd_gct'])
    output.date = pd.to_datetime(output.date, infer_datetime_format=True)

    return (output)

def create_cadence_plot(run_dates, metric, title_string):
    cmap = plt.cm.get_cmap('Set1')
    # moving_avg = pd.DataFrame({'dates':run_dates, 'ma':participant_data.mean_metric.rolling(window=7).mean()})
    # moving_avg.dates = pd.to_datetime(moving_avg.dates)
    # moving_avg = moving_avg.set_index(['dates'])
    # upsampled = moving_avg.resample('D')
    # interpolated = upsampled.interpolate(method='spline', order = 5)

    ##upsampled_dates = interpolated.index.values
    # upsampled_vals = interpolated.ma.values
    ylims = [min([np.nanmin(metric - 5), 175]), max([np.nanmax(metric + 5), 190])]
    xlims = [min(run_dates) - np.timedelta64(5, "D"), max(run_dates) + np.timedelta64(5, "D")]

    mondays = WeekdayLocator(MONDAY)

    months = MonthLocator(range(1, 13), bymonthday=1)
    monthsFmt = DateFormatter("%d %b %Y")

    fig, ax = plt.subplots()
    #ax.fill_between(xlims, 180, 190, facecolor=cmap(2), alpha=0.5)
    #ax.fill_between(xlims, 165, 180, facecolor=cmap(4), alpha=0.5)
    #ax.fill_between(xlims, 0, 165, facecolor=cmap(0), alpha=0.5)
    ax.axhline(180, color = 'black', alpha = 0.5, linestyle = '-.', zorder = 0)
    ax.bar(run_dates, metric, color = cmap(1), edgecolor = "white", linewidth=0.1)

    # ax.plot(upsampled_dates, upsampled_vals)
    # ax.axhline(np.mean(metric), color = cmap(3), alpha = 0.2)

    ax.xaxis.set_major_locator(mondays)
    ax.xaxis.set_major_formatter(monthsFmt)
    ax.xaxis.set_minor_locator(mondays)
    ax.set_ylim(ylims[0], ylims[1])
    ax.set_xlim(xlims[0], xlims[1])
    ax.set_title(title_string, fontsize=14, color='grey')
    fig.autofmt_xdate()
    ax.margins(0.01)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.set_size_inches(6, 3.54331)
    plt.xticks(fontsize=6)

    if not os.path.exists(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images'):
        os.makedirs(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images')
    plt.savefig(
         FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/candence.png',
        bbox_inches='tight', dpi=300)

    return (fig)

def create_breaking_plot(run_dates, metric, title_string):
    cmap = plt.cm.get_cmap('Set1')
    # moving_avg = pd.DataFrame({'dates':run_dates, 'ma':participant_data.mean_metric.rolling(window=7).mean()})
    # moving_avg.dates = pd.to_datetime(moving_avg.dates)
    # moving_avg = moving_avg.set_index(['dates'])
    # upsampled = moving_avg.resample('D')
    # interpolated = upsampled.interpolate(method='spline', order = 5)

    ##upsampled_dates = interpolated.index.values
    # upsampled_vals = interpolated.ma.values

    ylims = [0, max([np.nanmax(metric) + 0.1, 0.4])]
    xlims = [min(run_dates) - np.timedelta64(5, "D"), max(run_dates) + np.timedelta64(5, "D")]
    mondays = WeekdayLocator(MONDAY)

    months = MonthLocator(range(1, 13), bymonthday=1)
    monthsFmt = DateFormatter("%d %b %Y")

    fig, ax = plt.subplots()
    #ax.fill_between(xlims, 0, 0.4, facecolor=cmap(2), alpha=0.5)
    #ax.fill_between(xlims, 0.4, 0.7, facecolor=cmap(4), alpha=0.5)
    #ax.fill_between(xlims, 0.7, 10, facecolor=cmap(0), alpha=0.5)
    ax.axhline(0.4, color = 'black', alpha = 0.5, linestyle = '-.', zorder = 0)
    ax.bar(run_dates, metric, color = cmap(1), edgecolor = "white", linewidth=0.1)
    # ax.plot(upsampled_dates, upsampled_vals)
    # ax.axhline(np.mean(metric), color = cmap(3), alpha = 0.2)

    ax.xaxis.set_major_locator(mondays)
    ax.xaxis.set_major_formatter(monthsFmt)
    ax.xaxis.set_minor_locator(mondays)
    ax.set_ylim(ylims[0], ylims[1])
    ax.set_xlim(xlims[0], xlims[1])
    ax.set_title(title_string, fontsize=14, color='grey')
    fig.autofmt_xdate()
    ax.margins(0.01)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.set_size_inches(6, 3.54331)
    plt.xticks(fontsize=6)
    if not os.path.exists(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images'):
        os.makedirs(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images')
    plt.savefig(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/breaking.png',
        bbox_inches='tight', dpi=300)
    fig.autofmt_xdate()
    plt.show

    return (fig)

def create_bounce_plot(run_dates, metric, title_string):
    cmap = plt.cm.get_cmap('Set1')
    # moving_avg = pd.DataFrame({'dates':run_dates, 'ma':participant_data.mean_metric.rolling(window=7).mean()})
    # moving_avg.dates = pd.to_datetime(moving_avg.dates)
    # moving_avg = moving_avg.set_index(['dates'])
    # upsampled = moving_avg.resample('D')
    # interpolated = upsampled.interpolate(method='spline', order = 5)

    ##upsampled_dates = interpolated.index.values
    # upsampled_vals = interpolated.ma.values

    ylims = [0, max([np.nanmax(metric) + 0.5, 8.2])]
    xlims = [min(run_dates) - np.timedelta64(5, "D"), max(run_dates) + np.timedelta64(5, "D")]
    mondays = WeekdayLocator(MONDAY)

    months = MonthLocator(range(1, 13), bymonthday=1)
    monthsFmt = DateFormatter("%d %b %Y")

    fig, ax = plt.subplots()
    #ax.fill_between(xlims, 0, 8.2, facecolor=cmap(2), alpha=0.5)
    #ax.fill_between(xlims, 8.2, 12, facecolor=cmap(4), alpha=0.5)
    #ax.fill_between(xlims, 12, 100, facecolor=cmap(0), alpha=0.5)
    ax.axhline(8.2, color = 'black', alpha = 0.5, linestyle = '-.', zorder = 0)
    ax.bar(run_dates, metric, color = cmap(1), edgecolor = "white", linewidth=0.1)
    # ax.plot(upsampled_dates, upsampled_vals)
    # ax.axhline(np.mean(metric), color = cmap(3), alpha = 0.2)

    ax.xaxis.set_major_locator(mondays)
    ax.xaxis.set_major_formatter(monthsFmt)
    ax.xaxis.set_minor_locator(mondays)
    ax.set_ylim(ylims[0], ylims[1])
    ax.set_xlim(xlims[0], xlims[1])
    ax.set_title(title_string, fontsize=14, color='grey')
    fig.autofmt_xdate()
    ax.margins(0.01)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.set_size_inches(6, 3.54331)
    plt.xticks(fontsize=6)
    if not os.path.exists(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images'):
        os.makedirs(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images')
    plt.savefig(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/bounce.png',
        bbox_inches='tight', dpi=300)
    fig.autofmt_xdate()
    plt.show

    return (fig)

def create_rotation_plot(run_dates, metric, title_string):
    cmap = plt.cm.get_cmap('Set1')
    # moving_avg = pd.DataFrame({'dates':run_dates, 'ma':participant_data.mean_metric.rolling(window=7).mean()})
    # moving_avg.dates = pd.to_datetime(moving_avg.dates)
    # moving_avg = moving_avg.set_index(['dates'])
    # upsampled = moving_avg.resample('D')
    # interpolated = upsampled.interpolate(method='spline', order = 5)

    ##upsampled_dates = interpolated.index.values
    # upsampled_vals = interpolated.ma.values

    ylims = [0, max([np.nanmax(metric) + 2.5, 15])]
    xlims = [min(run_dates) - np.timedelta64(5, "D"), max(run_dates) + np.timedelta64(5, "D")]
    mondays = WeekdayLocator(MONDAY)

    months = MonthLocator(range(1, 13), bymonthday=1)
    monthsFmt = DateFormatter("%d %b %Y")

    fig, ax = plt.subplots()
    #ax.fill_between(xlims, 0, 15, facecolor=cmap(2), alpha=0.5)
    #ax.fill_between(xlims, 15, 18, facecolor=cmap(4), alpha=0.5)
    #ax.fill_between(xlims, 18, 100, facecolor=cmap(0), alpha=0.5)
    ax.axhline(15, color = 'black', alpha = 0.5, linestyle = '-.', zorder=0)
    ax.bar(run_dates, metric, color = cmap(1), edgecolor="white", linewidth=0.1)

    # ax.plot(upsampled_dates, upsampled_vals)
    # ax.axhline(np.mean(metric), color = cmap(3), alpha = 0.2)

    ax.xaxis.set_major_locator(mondays)
    ax.xaxis.set_major_formatter(monthsFmt)
    ax.xaxis.set_minor_locator(mondays)
    ax.set_ylim(ylims[0], ylims[1])
    ax.set_xlim(xlims[0], xlims[1])
    ax.set_title(title_string, fontsize=14, color='grey')
    fig.autofmt_xdate()
    ax.margins(0.01)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.set_size_inches(6, 3.54331)
    plt.xticks(fontsize=6)
    if not os.path.exists(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images'):
        os.makedirs(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images')
    plt.savefig(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/rotation.png',
        bbox_inches='tight', dpi=300)
    fig.autofmt_xdate()
    plt.show

    return (fig)

def create_drop_plot(run_dates, metric, title_string):
    cmap = plt.cm.get_cmap('Set1')
    # moving_avg = pd.DataFrame({'dates':run_dates, 'ma':participant_data.mean_metric.rolling(window=7).mean()})
    # moving_avg.dates = pd.to_datetime(moving_avg.dates)
    # moving_avg = moving_avg.set_index(['dates'])
    # upsampled = moving_avg.resample('D')
    # interpolated = upsampled.interpolate(method='spline', order = 5)

    ##upsampled_dates = interpolated.index.values
    # upsampled_vals = interpolated.ma.values

    ylims = [0, max([np.nanmax(metric) + 2.5, 15])]
    xlims = [min(run_dates) - np.timedelta64(5, "D"), max(run_dates) + np.timedelta64(5, "D")]
    mondays = WeekdayLocator(MONDAY)

    months = MonthLocator(range(1, 13), bymonthday=1)
    monthsFmt = DateFormatter("%d %b %Y")

    fig, ax = plt.subplots()
    #ax.fill_between(xlims, 0, 12, facecolor=cmap(2), alpha=0.5)
    #ax.fill_between(xlims, 12, 20, facecolor=cmap(4), alpha=0.5)
    #ax.fill_between(xlims, 20, 100, facecolor=cmap(0), alpha=0.5)
    ax.axhline(12, color='black', alpha=0.5, linestyle='-.', zorder=0)
    ax.bar(run_dates, metric, color=cmap(1), edgecolor="white", linewidth=0.1)

    # ax.plot(upsampled_dates, upsampled_vals)
    # ax.axhline(np.mean(metric), color = cmap(3), alpha = 0.2)

    ax.xaxis.set_major_locator(mondays)
    ax.xaxis.set_major_formatter(monthsFmt)
    ax.xaxis.set_minor_locator(mondays)
    ax.set_ylim(ylims[0], ylims[1])
    ax.set_xlim(xlims[0], xlims[1])
    ax.set_title(title_string, fontsize=14, color='grey')
    fig.autofmt_xdate()
    ax.margins(0.01)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.set_size_inches(6, 3.54331)
    plt.xticks(fontsize=6)
    if not os.path.exists(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images'):
        os.makedirs(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images')
    plt.savefig(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/drop.png',
        bbox_inches='tight', dpi=300)
    fig.autofmt_xdate()
    plt.show

    return (fig)

def search_paragraph_text(document, re_text):
    for i in range(0, len(document.paragraphs)):
        text = document.paragraphs[i].text
        if re.search(re_text, text):
            # print(i, 'match')
            return (i)

def create_word_doc(participant_data):
    print('Processing word doc....')
    doc = docx.Document('report_template.docx')

    # Update subject id

    doc.paragraphs[2].text = ""
    run = doc.paragraphs[2].add_run('Subject id: ').bold = True
    doc.paragraphs[2].add_run(PARTICIPANT_ID)

    # update Basic stats table table
    basic_stats = doc.tables[0]
    basic_stats.rows[1].cells[0].paragraphs[0].runs[0].text = str(participant_data.shape[0])

    run_dates = participant_data.date.values
    n_weeks = np.timedelta64(max(run_dates) - min(run_dates), "W") / np.timedelta64(1, 'W')  # 18 weeks
    rpw = round(len(run_dates) / n_weeks, 1)
    basic_stats.rows[1].cells[1].paragraphs[0].runs[0].text = str(
        rpw)
    # by number of weeks

    # typical duration
    duration = participant_data.time.values
    duration = duration[duration != 0]
    duration = duration[~np.isnan(duration)]
    typical_duration = np.percentile(duration, 50)
    m, s = divmod(typical_duration, 60)
    h, m = divmod(m, 60)
    basic_stats.rows[4].cells[0].paragraphs[0].runs[0].text = str("%d:%02d:%02d" % (h, m, s))

    # typical distance
    distance = participant_data.distance.values
    distance = distance[distance != 0]
    distance = distance[~np.isnan(distance)]
    typical_distance = np.percentile(distance, 50)
    basic_stats.rows[4].cells[1].paragraphs[0].runs[0].text = str(round(typical_distance / 1000, 2))

    # typical Speed
    speed = participant_data.mean_speed.values
    speed = speed[speed != 0]
    speed = speed[~np.isnan(speed)]
    typical_speed = np.percentile(speed, 50)
    basic_stats.rows[7].cells[0].paragraphs[0].runs[0].text = str(round(typical_speed * 3.6, 2))

    # typical elevation
    elevation = participant_data.elevation_gain.values
    elevation = elevation[elevation != 0]
    elevation = elevation[~np.isnan(elevation)]
    typical_elevation = np.percentile(elevation, 50)
    basic_stats.rows[7].cells[1].paragraphs[0].runs[0].text = str(round(typical_elevation, 2))
    doc.tables[0] = basic_stats

    ## Cadence
    # Update Cadence table
    cadence = participant_data.mean_cadence.values
    cadence = cadence[cadence != 0]
    cadence = cadence[~np.isnan(cadence)]
    typical_cadence = int(round(np.percentile(cadence, 50), 0))
    cadence_table = doc.tables[1]
    cadence_table.rows[1].cells[1].paragraphs[0].runs[0].text = str("{} steps per min".format(typical_cadence))

    # Add Cadence Figure
    p = search_paragraph_text(doc, r'Add Cadence Image Here')
    doc.paragraphs[p].runs[0].text = ""
    doc.paragraphs[p].runs[0].add_picture(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/candence.png',
        width=docx.shared.Inches(6), height=docx.shared.Inches(3.54331))

    ## Breaking
    # Update breaking table
    breaking = participant_data.mean_breaking.values
    breaking = breaking[breaking != 0]
    breaking = breaking[~np.isnan(breaking)]
    typical_breaking = round(np.percentile(breaking, 50), 1)
    breaking_table = doc.tables[2]
    breaking_table.rows[1].cells[1].paragraphs[0].runs[0].text = str("{} m/s".format(typical_breaking))

    # Add breakingFigure
    p = search_paragraph_text(doc, r'Add Breaking Image Here')
    doc.paragraphs[p].runs[0].text = ""
    doc.paragraphs[p].runs[0].add_picture(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/breaking.png',
        width=docx.shared.Inches(6), height=docx.shared.Inches(3.54331))

    ## Bounce
    # Update bounce table
    bounce = participant_data.mean_bounce.values
    bounce = bounce[bounce != 0]
    bounce = bounce[~np.isnan(bounce)]
    typical_bounce = round(np.percentile(bounce, 50), 1)
    bounce_table = doc.tables[3]
    bounce_table.rows[1].cells[1].paragraphs[0].runs[0].text = str("{} cm".format(typical_bounce))

    # Add bounce Figure
    p = search_paragraph_text(doc, r'Add Bounce Image Here')
    doc.paragraphs[p].runs[0].text = ""
    doc.paragraphs[p].runs[0].add_picture(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/bounce.png',
        width=docx.shared.Inches(6), height=docx.shared.Inches(3.54331))

    ## Rotation
    # Update rotation table
    rotation = participant_data.mean_rotation.values
    rotation = rotation[rotation != 0]
    rotation = rotation[~np.isnan(rotation)]
    typical_rotation = round(np.percentile(rotation, 50), 1)
    rotation_table = doc.tables[4]
    rotation_table.rows[1].cells[1].paragraphs[0].runs[0].text = str("{}\u00b0".format(typical_rotation))

    # Add rotation Figure
    p = search_paragraph_text(doc, r'Add Rotation Image Here')
    doc.paragraphs[p].runs[0].text = ""
    doc.paragraphs[p].runs[0].add_picture(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/rotation.png',
        width=docx.shared.Inches(6), height=docx.shared.Inches(3.54331))

    ## Drop
    # Update drop table
    drop = participant_data.mean_drop.values
    drop = drop[drop != 0]
    drop = drop[~np.isnan(drop)]
    typical_drop = round(np.percentile(drop, 50), 1)
    drop_table = doc.tables[5]
    drop_table.rows[1].cells[1].paragraphs[0].runs[0].text = str("{}\u00b0".format(typical_drop))

    # Add rotation Figure
    p = search_paragraph_text(doc, r'Add Drop Image Here')
    doc.paragraphs[p].runs[0].text = ""
    doc.paragraphs[p].runs[0].add_picture(
        FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp/images/drop.png',
        width=docx.shared.Inches(6), height=docx.shared.Inches(3.54331))

    if not os.path.exists(FILE_DIR + str(PARTICIPANT_ID) + '/reports'):
        os.makedirs(FILE_DIR + str(PARTICIPANT_ID) + '/reports')

    file_string = FILE_DIR + str(PARTICIPANT_ID) + '/reports/Biomechanics_report_' + PARTICIPANT_ID + '_'+ END_DATE +'.docx'

    doc.save(file_string)

    print('Word doc complete')
    print('Converting to pdf....')
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(file_string.replace('/', '\\'))

    pdf_string = FILE_DIR + '/' + str(PARTICIPANT_ID) + '/reports/Biomechanics_report_' + PARTICIPANT_ID + '_'+ END_DATE +'.pdf'
    doc.SaveAs(pdf_string.replace('/', '\\'), FileFormat=17)
    doc.Close()
    word.Quit()

    shutil.rmtree(FILE_DIR + '/' + str(PARTICIPANT_ID) + '/tmp') 
    print('Complete!!')


def main():
    file_list = get_dates(PARTICIPANT_ID)
    participant_data = get_participant_data(file_list)
    print('Generating images...')
    create_cadence_plot(participant_data.date.values, participant_data.mean_cadence, 'Cadence (SPM)')
    create_breaking_plot(participant_data.date.values, participant_data.mean_breaking, 'Breaking (m/s)')
    create_bounce_plot(participant_data.date.values, participant_data.mean_bounce, 'Bounce (cm)')
    create_rotation_plot(participant_data.date.values, participant_data.mean_rotation, 'Rotation (\u00b0)')
    create_drop_plot(participant_data.date.values, participant_data.mean_drop, 'Drop (\u00b0)')
    create_word_doc(participant_data)


if __name__ == "__main__":
    main()
